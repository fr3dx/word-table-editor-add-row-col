import logging
from typing import Optional, Tuple

from docx import Document
from docx.table import Table

from src.table_editor import insert_row_and_column

# Constants
DEFAULT_INPUT_FILE = "example_table.docx"
DEFAULT_OUTPUT_FILE = "document_row_column_modified.docx"
USER_PROMPTS = {
    'insert_row': "➕ Do you want to insert a new row? (y/n/q): ",
    'row_position': "➡️ At which row position? (1-{}, 'q' to skip): ",
    'insert_col': "➕ Do you want to insert a column? (y/n/q): ",
    'col_position': "➡️ At which column position? (1-{}, 'q' to skip): ",
}

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)


def get_user_input(prompt: str, valid_range: Optional[Tuple[int, int]] = None) -> Optional[int]:
    """
    Get validated user input.
    
    Args:
        prompt: The input prompt to display
        valid_range: Optional range of valid numbers (min, max)
    
    Returns:
        The validated input or None if user quits
    """
    while True:
        user_input = input(prompt).strip().lower()
        
        if user_input == 'q':
            return None
            
        if not user_input.isdigit():
            logger.warning("❗ Please enter a valid number or 'q' to quit.")
            continue
            
        num = int(user_input)
        
        if valid_range and not (valid_range[0] <= num <= valid_range[1]):
            logger.warning(f"❗ Please enter a number between {valid_range[0]} and {valid_range[1]}.")
            continue
            
        return num


def process_table(table: Table, table_idx: int) -> Tuple[Optional[int], Optional[int]]:
    """
    Process a single table to get user input for row/column insertion.
    
    Args:
        table: The table to process
        table_idx: Index of the table in the document
    
    Returns:
        Tuple of (row_position, column_position) or (None, None)
    """
    row_count = len(table.rows)
    col_count = len(table.columns)
    
    logger.info(f"\n📊 [Table {table_idx+1}] Rows: {row_count}, Columns: {col_count}")

    row_pos = None
    if input(USER_PROMPTS['insert_row']).strip().lower() == 'y':
        row_pos = get_user_input(
            USER_PROMPTS['row_position'].format(row_count + 1),
            (1, row_count + 1)
        )

    col_pos = None
    if input(USER_PROMPTS['insert_col']).strip().lower() == 'y':
        col_pos = get_user_input(
            USER_PROMPTS['col_position'].format(col_count + 1),
            (1, col_count + 1)
        )

    return row_pos, col_pos


def main() -> None:
    """Main function to process the Word document tables."""
    try:
        doc = Document(DEFAULT_INPUT_FILE)
        
        if not doc.tables:
            logger.warning("❗️No tables found in the document.")
            return

        logger.info(f"📄 The document contains {len(doc.tables)} table(s).")
        modified = False

        for idx, table in enumerate(doc.tables):
            row_pos, col_pos = process_table(table, idx)
            
            if row_pos is None and col_pos is None:
                logger.info("⏭️ Table unchanged.")
                continue
                
            if insert_row_and_column(doc, table, row_pos, col_pos):
                modified = True
                logger.info(
                    f"✅ Inserted: row {row_pos if row_pos else '-'}, "
                    f"column {col_pos if col_pos else '-'}"
                )

        if modified:
            doc.save(DEFAULT_OUTPUT_FILE)
            logger.info("\n💾 Document saved.")
        else:
            logger.info("\nℹ️ No changes made.")

    except FileNotFoundError:
        logger.error(f"❌ File not found: {DEFAULT_INPUT_FILE}")
    except Exception as e:
        logger.error(f"❌ An error occurred: {str(e)}")


if __name__ == "__main__":
    main()